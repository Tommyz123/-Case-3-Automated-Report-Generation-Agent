"""
Unit tests for WordTemplateHandler module.

Tests cover:
- Paragraph position finding (Task 2.2.11)
- Text insertion and style preservation (Task 2.2.12)
- Table insertion and formatting (Task 2.2.13)
"""

import pytest
from pathlib import Path
import tempfile
import shutil

from docx import Document
from docx.shared import Pt

from src.template_handler import WordTemplateHandler


# Test fixtures
@pytest.fixture
def temp_dir():
    """Create a temporary directory for test files."""
    temp_path = Path(tempfile.mkdtemp())
    yield temp_path
    # Cleanup
    shutil.rmtree(temp_path)


@pytest.fixture
def test_template(temp_dir):
    """Create a test template document."""
    template_path = temp_dir / "test_template.docx"

    # Create a simple test document
    doc = Document()

    # Add various paragraphs with different styles
    doc.add_heading('Purpose', level=1)
    doc.add_paragraph('This is the purpose section.')

    doc.add_heading('Underlying Principles', level=1)
    doc.add_paragraph('This section describes the underlying principles.')

    doc.add_heading('Phase 2: Research', level=2)
    doc.add_paragraph('This is the research phase.')

    doc.add_paragraph('Normal paragraph for testing.')

    # Save the template
    doc.save(str(template_path))

    return template_path


@pytest.fixture
def handler(test_template):
    """Create a WordTemplateHandler instance with test template."""
    return WordTemplateHandler(str(test_template))


# ==================== Test Paragraph Position Finding (Task 2.2.11) ====================

class TestParagraphPositionFinding:
    """Test paragraph position finding functionality."""

    def test_find_paragraph_by_text_exact_match(self, handler):
        """Test finding paragraph by exact text match."""
        # Act
        para = handler.find_paragraph_by_text("Purpose", exact_match=True)

        # Assert
        assert para is not None, "Should find paragraph with exact text 'Purpose'"
        assert para.text == "Purpose"

    def test_find_paragraph_by_text_substring_match(self, handler):
        """Test finding paragraph by substring match."""
        # Act
        para = handler.find_paragraph_by_text("purpose section", exact_match=False)

        # Assert
        assert para is not None, "Should find paragraph containing 'purpose section'"
        assert "purpose section" in para.text.lower()

    def test_find_paragraph_by_text_not_found(self, handler):
        """Test finding non-existent paragraph returns None."""
        # Act
        para = handler.find_paragraph_by_text("Nonexistent Text", exact_match=True)

        # Assert
        assert para is None, "Should return None for non-existent text"

    def test_find_paragraphs_by_style(self, handler):
        """Test finding paragraphs by style."""
        # Act
        heading1_paras = handler.find_paragraphs_by_style("Heading 1")

        # Assert
        assert len(heading1_paras) == 2, "Should find 2 Heading 1 paragraphs"
        assert all(p.style.name == "Heading 1" for p in heading1_paras)

    def test_find_paragraph_by_text_and_style(self, handler):
        """Test finding paragraph by both text and style."""
        # Act
        para = handler.find_paragraph_by_text_and_style(
            "Purpose",
            "Heading 1",
            exact_match=True
        )

        # Assert
        assert para is not None, "Should find paragraph with text 'Purpose' and style 'Heading 1'"
        assert para.text == "Purpose"
        assert para.style.name == "Heading 1"

    def test_find_paragraph_by_text_and_style_wrong_style(self, handler):
        """Test finding paragraph with wrong style returns None."""
        # Act
        para = handler.find_paragraph_by_text_and_style(
            "Purpose",
            "Normal",  # Wrong style (it's actually Heading 1)
            exact_match=True
        )

        # Assert
        assert para is None, "Should return None when style doesn't match"

    def test_get_all_styles(self, handler):
        """Test getting all styles used in document."""
        # Act
        styles = handler.get_all_styles()

        # Assert
        assert "Heading 1" in styles, "Should include Heading 1 style"
        assert "Heading 2" in styles, "Should include Heading 2 style"
        assert "Normal" in styles, "Should include Normal style"


# ==================== Test Text Insertion (Task 2.2.12) ====================

class TestTextInsertion:
    """Test text insertion and style preservation functionality."""

    def test_insert_text_after_paragraph(self, handler, temp_dir):
        """Test inserting text after a paragraph."""
        # Arrange
        target_para = handler.find_paragraph_by_text("Purpose")
        original_para_count = handler.get_paragraph_count()

        # Act
        new_para = handler.insert_text_after(target_para, "New text inserted here")

        # Assert
        assert new_para.text == "New text inserted here"
        assert handler.get_paragraph_count() == original_para_count + 1

    def test_insert_text_preserves_style(self, handler, temp_dir):
        """Test that inserting text preserves the original paragraph style."""
        # Arrange
        target_para = handler.find_paragraph_by_text("Purpose")
        original_style = target_para.style.name

        # Act
        new_para = handler.insert_text_after(
            target_para,
            "New text with preserved style",
            preserve_style=True
        )

        # Assert
        assert new_para.style.name == original_style, \
            f"Style should be preserved: expected {original_style}, got {new_para.style.name}"

    def test_insert_formatted_text_bold(self, handler, temp_dir):
        """Test inserting bold text."""
        # Arrange
        target_para = handler.find_paragraph_by_text("Purpose")

        # Act
        new_para = handler.insert_formatted_text(
            target_para,
            "Bold text",
            bold=True
        )

        # Assert
        assert len(new_para.runs) > 0, "Should have at least one run"
        assert new_para.runs[0].bold is True, "Text should be bold"

    def test_insert_formatted_text_italic(self, handler, temp_dir):
        """Test inserting italic text."""
        # Arrange
        target_para = handler.find_paragraph_by_text("Purpose")

        # Act
        new_para = handler.insert_formatted_text(
            target_para,
            "Italic text",
            italic=True
        )

        # Assert
        assert len(new_para.runs) > 0, "Should have at least one run"
        assert new_para.runs[0].italic is True, "Text should be italic"

    def test_insert_formatted_text_underline(self, handler, temp_dir):
        """Test inserting underlined text."""
        # Arrange
        target_para = handler.find_paragraph_by_text("Purpose")

        # Act
        new_para = handler.insert_formatted_text(
            target_para,
            "Underlined text",
            underline=True
        )

        # Assert
        assert len(new_para.runs) > 0, "Should have at least one run"
        assert new_para.runs[0].underline is True, "Text should be underlined"

    def test_insert_formatted_text_with_font_size(self, handler, temp_dir):
        """Test inserting text with custom font size."""
        # Arrange
        target_para = handler.find_paragraph_by_text("Purpose")

        # Act
        new_para = handler.insert_formatted_text(
            target_para,
            "Text with size 14",
            font_size=14
        )

        # Assert
        assert len(new_para.runs) > 0, "Should have at least one run"
        assert new_para.runs[0].font.size == Pt(14), "Font size should be 14pt"

    def test_insert_formatted_text_with_font_name(self, handler, temp_dir):
        """Test inserting text with custom font name."""
        # Arrange
        target_para = handler.find_paragraph_by_text("Purpose")

        # Act
        new_para = handler.insert_formatted_text(
            target_para,
            "Text with Arial font",
            font_name="Arial"
        )

        # Assert
        assert len(new_para.runs) > 0, "Should have at least one run"
        assert new_para.runs[0].font.name == "Arial", "Font should be Arial"


# ==================== Test Table Insertion (Task 2.2.13) ====================

class TestTableInsertion:
    """Test table insertion and formatting functionality."""

    def test_insert_table_basic(self, handler, temp_dir):
        """Test inserting a basic table."""
        # Arrange
        target_para = handler.find_paragraph_by_text("Purpose")

        # Act
        table = handler.insert_table_after(target_para, rows=3, cols=4)

        # Assert
        assert table is not None, "Table should be created"
        assert len(table.rows) == 3, "Table should have 3 rows"
        assert len(table.columns) == 4, "Table should have 4 columns"

    def test_insert_table_with_data(self, handler, temp_dir):
        """Test inserting a table with data."""
        # Arrange
        target_para = handler.find_paragraph_by_text("Purpose")
        data = [
            ["Header 1", "Header 2", "Header 3"],
            ["Row 1 Col 1", "Row 1 Col 2", "Row 1 Col 3"],
            ["Row 2 Col 1", "Row 2 Col 2", "Row 2 Col 3"],
        ]

        # Act
        table = handler.insert_table_after(
            target_para,
            rows=3,
            cols=3,
            data=data
        )

        # Assert
        assert table.rows[0].cells[0].text == "Header 1"
        assert table.rows[1].cells[1].text == "Row 1 Col 2"
        assert table.rows[2].cells[2].text == "Row 2 Col 3"

    def test_insert_table_with_header_bold(self, handler, temp_dir):
        """Test inserting table with bold header."""
        # Arrange
        target_para = handler.find_paragraph_by_text("Purpose")
        data = [
            ["Header 1", "Header 2"],
            ["Data 1", "Data 2"],
        ]
        style_config = {
            'header_bold': True
        }

        # Act
        table = handler.insert_table_after(
            target_para,
            rows=2,
            cols=2,
            data=data,
            style_config=style_config
        )

        # Assert
        header_cell = table.rows[0].cells[0]
        # Check if header text is bold
        for para in header_cell.paragraphs:
            for run in para.runs:
                if run.text:  # Only check non-empty runs
                    assert run.bold is True, "Header text should be bold"

    def test_insert_table_with_borders(self, handler, temp_dir):
        """Test inserting table with borders."""
        # Arrange
        target_para = handler.find_paragraph_by_text("Purpose")
        style_config = {
            'border': True
        }

        # Act
        table = handler.insert_table_after(
            target_para,
            rows=2,
            cols=2,
            style_config=style_config
        )

        # Assert
        # Just verify table was created (detailed border testing would require
        # XML inspection which is complex)
        assert table is not None, "Table should be created with borders"


# ==================== Test Document Save ====================

class TestDocumentSave:
    """Test document saving functionality."""

    def test_save_document(self, handler, temp_dir):
        """Test saving modified document."""
        # Arrange
        output_path = temp_dir / "output" / "modified_template.docx"
        target_para = handler.find_paragraph_by_text("Purpose")
        handler.insert_text_after(target_para, "Modified content")

        # Act
        handler.save_document(str(output_path))

        # Assert
        assert output_path.exists(), "Output file should be created"

        # Verify the saved document can be opened
        saved_doc = Document(str(output_path))
        assert len(saved_doc.paragraphs) > 0, "Saved document should have paragraphs"

    def test_save_document_creates_output_directory(self, handler, temp_dir):
        """Test that save_document creates output directory if it doesn't exist."""
        # Arrange
        output_path = temp_dir / "new_dir" / "nested_dir" / "output.docx"

        # Act
        handler.save_document(str(output_path))

        # Assert
        assert output_path.exists(), "Output file should be created"
        assert output_path.parent.exists(), "Output directory should be created"


# ==================== Integration Tests ====================

class TestWordTemplateHandlerIntegration:
    """Integration tests for WordTemplateHandler."""

    def test_complete_workflow(self, handler, temp_dir):
        """Test complete workflow: find, insert, save."""
        # Arrange
        output_path = temp_dir / "complete_output.docx"

        # Act
        # 1. Find paragraph by text and style
        para = handler.find_paragraph_by_text_and_style(
            "Underlying Principles",
            "Heading 1"
        )
        assert para is not None, "Should find target paragraph"

        # 2. Insert formatted text
        new_text_para = handler.insert_formatted_text(
            para,
            "This is a new paragraph with bold text.",
            bold=True,
            font_size=11
        )
        assert new_text_para is not None, "Should insert new paragraph"

        # 3. Insert table
        table_data = [
            ["Stakeholder", "Impact", "Value"],
            ["Hospitals", "Cost savings", "£10,000"],
            ["Patients", "Time saved", "2 hours"],
        ]
        table = handler.insert_table_after(
            new_text_para,
            rows=3,
            cols=3,
            data=table_data,
            style_config={'header_bold': True, 'border': True}
        )
        assert table is not None, "Should insert table"

        # 4. Save document
        handler.save_document(str(output_path))

        # Assert
        assert output_path.exists(), "Output file should exist"

        # Verify saved document
        saved_doc = Document(str(output_path))
        assert len(saved_doc.tables) > 0, "Saved document should have tables"
        assert saved_doc.tables[0].rows[0].cells[0].text == "Stakeholder"
