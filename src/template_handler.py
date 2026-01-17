"""
Word Template Handler Module for Case 3 Automation Agent.

This module handles Word document template operations, including:
- Loading Word templates
- Finding paragraph positions (by text or style)
- Inserting text content (preserving original document styles)
- Inserting formatted text (bold, italic, underline)
- Inserting tables
- Saving modified documents

根据 PROJECT_PLAN.md Phase 2.2 实现。
"""

import logging
from pathlib import Path
from typing import List, Optional, Dict, Any

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph
from docx.table import Table, _Cell

# Configure logging
logger = logging.getLogger(__name__)


class WordTemplateHandler:
    """
    Handles Word template operations.

    Responsibilities:
    - Load Word template documents
    - Find paragraph positions by text matching or style matching
    - Insert text content while preserving original document styles
    - Insert formatted text (bold, italic, underline)
    - Insert tables with custom styling
    - Save modified documents

    Usage:
        handler = WordTemplateHandler("template.docx")
        para = handler.find_paragraph_by_text("Chapter 1")
        handler.insert_text_after(para, "New content here")
        handler.save_document("output.docx")
    """

    def __init__(self, template_path: str):
        """
        Initialize WordTemplateHandler.

        Args:
            template_path: Path to the Word template file

        Raises:
            FileNotFoundError: If template file doesn't exist
        """
        self.template_path = Path(template_path)

        if not self.template_path.exists():
            raise FileNotFoundError(f"Template file not found: {self.template_path}")

        logger.info(f"Loading Word template: {self.template_path}")
        self.document = Document(str(self.template_path))
        logger.info(f"Template loaded successfully. Contains {len(self.document.paragraphs)} paragraphs")

    # ==================== Paragraph Position Finding ====================

    def find_paragraph_by_text(
        self,
        text: str,
        exact_match: bool = True
    ) -> Optional[Paragraph]:
        """
        Find a paragraph by text content.

        Args:
            text: Text to search for
            exact_match: If True, match exact text; if False, match substring

        Returns:
            Paragraph object if found, None otherwise
        """
        logger.debug(f"Searching for paragraph with text: '{text}' (exact_match={exact_match})")

        for para in self.document.paragraphs:
            para_text = para.text.strip()

            if exact_match:
                if para_text == text:
                    logger.debug(f"Found exact match: '{para_text}'")
                    return para
            else:
                if text in para_text:
                    logger.debug(f"Found substring match: '{para_text}'")
                    return para

        logger.warning(f"Paragraph not found with text: '{text}'")
        return None

    def find_paragraphs_by_style(self, style: str) -> List[Paragraph]:
        """
        Find all paragraphs with a specific style.

        Args:
            style: Style name (e.g., "Heading 1", "Heading 2", "Normal")

        Returns:
            List of Paragraph objects matching the style
        """
        logger.debug(f"Searching for paragraphs with style: '{style}'")

        matching_paragraphs = []

        for para in self.document.paragraphs:
            if para.style and para.style.name == style:
                matching_paragraphs.append(para)

        logger.debug(f"Found {len(matching_paragraphs)} paragraphs with style '{style}'")
        return matching_paragraphs

    def find_paragraph_by_text_and_style(
        self,
        text: str,
        style: str,
        exact_match: bool = True
    ) -> Optional[Paragraph]:
        """
        Find a paragraph by both text content and style.

        This is more precise than searching by text alone.

        Args:
            text: Text to search for
            style: Style name (e.g., "Heading 1")
            exact_match: If True, match exact text; if False, match substring

        Returns:
            Paragraph object if found, None otherwise
        """
        logger.debug(f"Searching for paragraph with text '{text}' and style '{style}'")

        for para in self.document.paragraphs:
            para_text = para.text.strip()
            para_style = para.style.name if para.style else None

            # Check style match
            if para_style != style:
                continue

            # Check text match
            if exact_match:
                if para_text == text:
                    logger.debug(f"Found match: '{para_text}' with style '{para_style}'")
                    return para
            else:
                if text in para_text:
                    logger.debug(f"Found match: '{para_text}' with style '{para_style}'")
                    return para

        logger.warning(f"Paragraph not found with text '{text}' and style '{style}'")
        return None

    # ==================== Text Insertion ====================

    def insert_text_after(
        self,
        position: Paragraph,
        text: str,
        preserve_style: bool = True
    ) -> Paragraph:
        """
        Insert text content after a specified paragraph.

        Args:
            position: Paragraph to insert after
            text: Text content to insert
            preserve_style: If True, copy the style from position paragraph

        Returns:
            The newly created paragraph
        """
        logger.debug(f"Inserting text after paragraph: '{position.text[:50]}...'")

        # Create new paragraph using OxmlElement
        new_para_element = OxmlElement('w:p')
        position._element.addnext(new_para_element)

        # Create a Paragraph object from the element
        new_para = Paragraph(new_para_element, position._parent)

        # Add text to the new paragraph
        if text:
            new_para.add_run(text)

        # Preserve style if requested
        if preserve_style and position.style:
            new_para.style = position.style
            logger.debug(f"Applied style: {position.style.name}")

        return new_para

    def insert_formatted_text(
        self,
        position: Paragraph,
        text: str,
        bold: bool = False,
        italic: bool = False,
        underline: bool = False,
        font_size: Optional[int] = None,
        font_name: Optional[str] = None
    ) -> Paragraph:
        """
        Insert formatted text after a specified paragraph.

        Args:
            position: Paragraph to insert after
            text: Text content to insert
            bold: Apply bold formatting
            italic: Apply italic formatting
            underline: Apply underline formatting
            font_size: Font size in points (e.g., 11)
            font_name: Font name (e.g., "Arial")

        Returns:
            The newly created paragraph
        """
        logger.debug(f"Inserting formatted text: bold={bold}, italic={italic}, underline={underline}")

        # Insert paragraph
        new_para = self.insert_text_after(position, text, preserve_style=False)

        # Apply formatting to all runs in the paragraph
        for run in new_para.runs:
            if bold:
                run.bold = True
            if italic:
                run.italic = True
            if underline:
                run.underline = True
            if font_size:
                run.font.size = Pt(font_size)
            if font_name:
                run.font.name = font_name

        return new_para

    # ==================== Table Insertion ====================

    def insert_table_after(
        self,
        position: Paragraph,
        rows: int,
        cols: int,
        data: Optional[List[List[str]]] = None,
        style_config: Optional[Dict[str, Any]] = None
    ) -> Table:
        """
        Insert a table after a specified paragraph.

        Args:
            position: Paragraph to insert after
            rows: Number of rows
            cols: Number of columns
            data: Optional table data (list of rows, each row is a list of cell values)
            style_config: Optional styling configuration dict with keys:
                - header_bold (bool): Make header row bold
                - header_background (str): Header background color (hex, e.g., "D3D3D3")
                - border (bool): Add borders to table

        Returns:
            The newly created table
        """
        logger.debug(f"Inserting table: {rows}x{cols}")

        # Create table
        table = self.document.add_table(rows=rows, cols=cols)

        # Move table to correct position (after the specified paragraph)
        position._element.addnext(table._element)

        # Fill data if provided
        if data:
            self._fill_table_data(table, data)

        # Apply styling if provided
        if style_config:
            self._apply_table_style(table, style_config)

        logger.debug(f"Table inserted successfully: {rows}x{cols}")
        return table

    def _fill_table_data(self, table: Table, data: List[List[str]]) -> None:
        """Fill table with data."""
        for row_idx, row_data in enumerate(data):
            if row_idx >= len(table.rows):
                # Add new row if needed
                table.add_row()

            row = table.rows[row_idx]

            for col_idx, cell_value in enumerate(row_data):
                if col_idx < len(row.cells):
                    row.cells[col_idx].text = str(cell_value) if cell_value is not None else ""

    def _apply_table_style(self, table: Table, style_config: Dict[str, Any]) -> None:
        """Apply styling to table."""
        # Apply header formatting
        if style_config.get('header_bold', False):
            header_row = table.rows[0]
            for cell in header_row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.bold = True

        # Apply header background color
        if 'header_background' in style_config:
            header_color = style_config['header_background']
            self._set_row_background_color(table.rows[0], header_color)

        # Apply borders
        if style_config.get('border', False):
            self._set_table_borders(table)

    def _set_row_background_color(self, row, color_hex: str) -> None:
        """Set background color for a table row."""
        for cell in row.cells:
            # Access cell shading element
            shading_elm = OxmlElement('w:shd')
            shading_elm.set(qn('w:fill'), color_hex)
            cell._element.get_or_add_tcPr().append(shading_elm)

    def _set_table_borders(self, table: Table) -> None:
        """Add borders to table."""
        # This is simplified - in production, you might want more control
        tbl = table._element
        tblPr = tbl.tblPr
        if tblPr is None:
            tblPr = OxmlElement('w:tblPr')
            tbl.insert(0, tblPr)

        # Add table borders element
        tblBorders = OxmlElement('w:tblBorders')
        for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), '4')  # Border size
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), '000000')  # Black
            tblBorders.append(border)

        tblPr.append(tblBorders)

    # ==================== Helper Methods ====================

    def _get_paragraph_index(self, paragraph: Paragraph) -> Optional[int]:
        """Get the index of a paragraph in the document."""
        try:
            return self.document.paragraphs.index(paragraph)
        except ValueError:
            return None

    def _insert_paragraph_at_index(self, index: int, text: str = "") -> Paragraph:
        """
        Insert a new paragraph at a specific index.

        Note: This uses internal XML API as python-docx doesn't provide
        a public API for inserting at specific positions.
        """
        # Get the paragraph element at the target position
        target_para = self.document.paragraphs[index] if index < len(self.document.paragraphs) else None

        # Create new paragraph
        new_para = self.document.add_paragraph(text)

        # Move it to the correct position
        if target_para:
            target_para._element.addprevious(new_para._element)
        # If index is at the end, paragraph is already in the right place

        return new_para

    # ==================== Document Save ====================

    def save_document(self, output_path: str) -> None:
        """
        Save the modified document to a file.

        Args:
            output_path: Path to save the document
        """
        output_file = Path(output_path)

        # Create output directory if it doesn't exist
        output_file.parent.mkdir(parents=True, exist_ok=True)

        logger.info(f"Saving document to: {output_file}")
        self.document.save(str(output_file))
        logger.info("Document saved successfully")

    # ==================== Utility Methods ====================

    def get_paragraph_count(self) -> int:
        """Get the total number of paragraphs in the document."""
        return len(self.document.paragraphs)

    def get_all_styles(self) -> List[str]:
        """Get a list of all styles used in the document."""
        styles = set()
        for para in self.document.paragraphs:
            if para.style:
                styles.add(para.style.name)
        return sorted(list(styles))
